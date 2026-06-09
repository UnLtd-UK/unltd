"""
Generate the UnLtd Hub Internal Engagement Plan as a Word document.
Run: python3 scripts/generate-engagement-plan.py
Output: UnLtd Hub — Internal Engagement Plan.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colours ───────────────────────────────────────────────────────────────────
UNLTD_RED   = RGBColor(0xE0, 0x2B, 0x2B)  # UnLtd primary red (approximation)
DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x2E)
MID_GREY    = RGBColor(0x6B, 0x7B, 0x8D)
LIGHT_BG    = RGBColor(0xF5, 0xF5, 0xF5)
TABLE_HEAD  = RGBColor(0x1A, 0x1A, 0x2E)
TABLE_ALT   = RGBColor(0xF9, 0xF9, 0xF9)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
AMBER       = RGBColor(0xFF, 0x8C, 0x00)
GREEN       = RGBColor(0x2E, 0x7D, 0x32)

def rgb_to_hex(color: RGBColor) -> str:
    # RGBColor is a tuple subclass: (r, g, b)
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

def set_cell_bg(cell, hex_color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_to_hex(hex_color))
    tcPr.append(shd)

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, enabled in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        border = OxmlElement(f"w:{side}")
        if enabled:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)
        else:
            border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph_shading(paragraph, hex_str: str):
    """Add background shading to a paragraph (for callout boxes)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    pPr.append(shd)

def add_run(paragraph, text, bold=False, italic=False,
            size=11, color=None, font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color=DARK_TEXT):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return p

def body(doc, text, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, size=11, color=DARK_TEXT)
    return p

def bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        add_run(p, bold_prefix + " ", bold=True, size=11, color=DARK_TEXT)
        add_run(p, text, size=11, color=DARK_TEXT)
    else:
        add_run(p, text, size=11, color=DARK_TEXT)
    return p

def callout(doc, text, bg_hex="EBF4FF", prefix="ℹ️  "):
    """Shaded callout / info box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    add_paragraph_shading(p, bg_hex)
    add_run(p, prefix + text, size=10, italic=True, color=MID_GREY)
    return p

def add_table_header_row(table, headers, col_widths_cm=None):
    row = table.rows[0]
    for i, (cell, header) in enumerate(zip(row.cells, headers)):
        set_cell_bg(cell, TABLE_HEAD)
        set_cell_borders(cell, color="1A1A2E")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, header, bold=True, size=10, color=WHITE)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            table.columns[i].width = Cm(w)

def add_data_row(table, values, alt=False):
    row = table.add_row()
    bg = TABLE_ALT if alt else WHITE
    for cell, value in zip(row.cells, values):
        set_cell_bg(cell, bg)
        set_cell_borders(cell, color="DDDDDD")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(value, tuple):
            text, bold = value
            add_run(p, text, bold=bold, size=10, color=DARK_TEXT)
        else:
            add_run(p, str(value), size=10, color=DARK_TEXT)

def spacer(doc, height_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(height_pt)
    run = p.add_run("")
    run.font.size = Pt(height_pt)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK_TEXT


# ─────────────────────────────────────────────────────────────────────────────
# TITLE BLOCK
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_run(p, "UnLtd Hub", bold=True, size=26, color=UNLTD_RED)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(4)
add_run(p2, "Internal Engagement Plan", bold=False, size=18, color=DARK_TEXT)

# Metadata row
meta = doc.add_table(rows=1, cols=4)
meta.style = "Table Grid"
meta_data = [
    ("Authors", "Tom Sheppard, Joel Attar"),
    ("Status", "Working Draft — For Review"),
    ("Date", "May 2026"),
    ("Target Internal Launch", "Early July 2026"),
]
for i, (cell, (label, value)) in enumerate(zip(meta.rows[0].cells, meta_data)):
    set_cell_bg(cell, LIGHT_BG)
    set_cell_borders(cell, color="CCCCCC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, label + "\n", bold=True, size=9, color=MID_GREY)
    add_run(p, value, bold=False, size=10, color=DARK_TEXT)

doc.add_paragraph()  # spacer


# ─────────────────────────────────────────────────────────────────────────────
# 1. PURPOSE OF THIS DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "1.  Purpose of This Document", level=1, color=UNLTD_RED)

body(doc,
     "This document sets out the internal engagement plan for the launch of UnLtd Hub — "
     "our new centralised knowledge platform built on GitBook. It is intended to ensure "
     "that all UnLtd colleagues, particularly those in the Social Entrepreneur Support (SES) "
     "Directorate and Delivery & Investment, are informed, prepared, and confident ahead of "
     "the public launch in September 2026.")

body(doc,
     "The plan covers: what UnLtd Hub is and why we're building it; how it will be used "
     "internally and externally; the rollout timeline; what each team needs to do; and the "
     "decisions and actions required to keep the project on track.")


# ─────────────────────────────────────────────────────────────────────────────
# 2. WHAT IS UNLTD HUB?
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "2.  What Is UnLtd Hub?", level=1, color=UNLTD_RED)

body(doc,
     "UnLtd Hub is a single, always up-to-date destination for social entrepreneurs "
     "throughout their entire journey with UnLtd — from first discovering who we are, "
     "through applying for an award, to their time as an award winner.")

body(doc,
     "It brings together resources that have previously been scattered across email, "
     "SharePoint, individual staff devices, and Canva. Rather than sending award winners "
     "a static welcome pack, the Hub provides one coherent, navigable space that can "
     "grow and improve over time.")

body(doc, "At launch it will be structured around three core areas:")

bullet(doc, "Business fundamentals and publicly available resources — accessible to anyone.",
       bold_prefix="Learn:")
bullet(doc, "Information about UnLtd awards and how to apply — accessible to anyone.",
       bold_prefix="About Awards:")
bullet(doc, "Onboarding materials, support resources, and cohort-specific content — "
            "accessible only to current award winners.",
       bold_prefix="Award Winner:")

spacer(doc, 8)
callout(doc,
        "Note: The About Awards section is subject to approval from Liam and Nicola in the "
        "context of an upcoming website rebrand. The content and structure of this section "
        "may change before the public launch.",
        bg_hex="FFF8E1", prefix="⚠️  ")


# ─────────────────────────────────────────────────────────────────────────────
# 3. WHY DO WE NEED IT?
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.  Why Do We Need It?", level=1, color=UNLTD_RED)

body(doc,
     "Our current approach to managing and sharing Award Winner resources is fragmented "
     "and inconsistent. The core problems are:")

problems = [
    ("No single source of truth.", "Resources are spread across SharePoint, Canva, individual "
     "staff desktops, and email. Award winners have no central place to go."),
    ("Siloed knowledge.", "Each SESM maintains their own informal resource collection. "
     "There is no complete picture of what exists across the team, and institutional knowledge "
     "is not shared."),
    ("No version control.", "Resources distributed as downloaded files quickly become "
     "outdated. Award winners may be working from old materials without knowing it."),
    ("No accessibility or quality standards.", "Resources vary significantly in quality, "
     "accessibility, and authorial style."),
    ("No analytics or feedback loops.", "We have no way of knowing which resources are "
     "being used or found valuable."),
    ("No access controls.", "Once a file is shared, we have no ability to restrict or manage "
     "who can view it."),
    ("A beta hub that was never launched.", "A resource hub was built on unltd.org.uk "
     "but never went live — the foundations already exist."),
]
for bold, text in problems:
    bullet(doc, text, bold_prefix=bold)


# ─────────────────────────────────────────────────────────────────────────────
# 4. WHAT IS GITBOOK AND WHY DID WE CHOOSE IT?
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "4.  What Is GitBook and Why Did We Choose It?", level=1, color=UNLTD_RED)

body(doc,
     "GitBook is a modern knowledge management and documentation platform. It lets teams "
     "create, organise, and publish structured content that is always online, always "
     "up-to-date, and accessible from any device. It is widely used by organisations that "
     "need a clean, professional space to share knowledge — both internally and publicly.")

body(doc, "We chose GitBook because it gives us:")

gitbook_reasons = [
    ("A single, always up-to-date source of truth.", "No more outdated files sent by email."),
    ("Structured, accessible content.", "Consistent formatting, style, and accessibility "
     "standards across everything we publish."),
    ("Built-in analytics.", "We can see which pages are being read, how often, and "
     "where people drop off — so we can improve the content over time."),
    ("Access permission controls.", "Award Winner content is gated behind access controls "
     "so only current award winners can see it."),
    ("Clear user roles.", "Editors, reviewers, and commenters have defined responsibilities "
     "and appropriate permissions — building accountability into the process."),
    ("A platform for future AI integration.", "GitBook has built-in AI capabilities that "
     "we can activate in a later phase to allow award winners to interact conversationally "
     "with content."),
    ("Free for our needs.", "The free tier covers all requirements for this phase."),
]
for bold, text in gitbook_reasons:
    bullet(doc, text, bold_prefix=bold)


# ─────────────────────────────────────────────────────────────────────────────
# 5. HOW WILL IT BE USED?
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "5.  How Will It Be Used?", level=1, color=UNLTD_RED)

heading(doc, "Internally — SES Directorate and Delivery & Investment", level=2, color=DARK_TEXT)
body(doc,
     "Support managers and UJSMs will use the Hub as a reference point and resource "
     "library when working with award winners. Rather than maintaining separate documents "
     "or searching through emails, they will direct award winners to specific pages within "
     "the Hub.")

body(doc,
     "The team will also play an active role in shaping the Hub, particularly in the early "
     "months. Support managers and UJSMs are best placed to identify gaps, flag "
     "inaccuracies, and suggest improvements based on their day-to-day experience of what "
     "social entrepreneurs actually need.")

heading(doc, "Externally — Award Winners", level=2, color=DARK_TEXT)
body(doc,
     "From July onwards, when Fern sends the welcome communication to a new cohort of "
     "award winners, the Hub link will replace the current welcome pack. Award winners "
     "will be directed to the Hub as their starting point for getting oriented, accessing "
     "their onboarding materials, and finding ongoing support resources.")

heading(doc, "Externally — Social Entrepreneurs (September onwards)", level=2, color=DARK_TEXT)
body(doc,
     "From September, the Hub will be made public. The Learn and About Awards sections "
     "will be available to anyone. The Award Winner section will remain behind access "
     "controls, visible only to current award winners.")

body(doc,
     "This makes the Hub the destination for the full social entrepreneur user journey: "
     "someone might first encounter it as a useful resource for signposting, come back to "
     "learn about UnLtd and apply, and then return as an award winner to access their "
     "cohort materials.")


# ─────────────────────────────────────────────────────────────────────────────
# 6. CONTENT: WHAT IS IN THE HUB?
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "6.  Content: What Is in the Hub?", level=1, color=UNLTD_RED)

callout(doc,
        "Migration is not one-to-one. Some resources will be merged, renamed, or newly "
        "created to ensure navigation is easy and content is consistent. The migration "
        "process has already surfaced mistakes and discrepancies in existing resources "
        "which are being corrected.",
        bg_hex="EBF4FF", prefix="ℹ️  ")

# ── Award Winner space ────────────────────────────────────────────────────────
heading(doc, "Award Winner Space", level=2, color=DARK_TEXT)
body(doc, "Gated — accessible only to current award winners.")

aw_headers = ["Page", "Status", "Consulted on first draft", "Editors", "Reviewer", "Commenters"]
aw_rows = [
    ("Welcome",                      "In progress",  "Fern",                       "SES Directorate", "User Journey Support Lead (Fern)", "All of UnLtd"),
    ("Your Support Manager",         "In progress",  "Tim, Fern",                  "SES Directorate", "",                                "All of UnLtd"),
    ("Get business mentoring",        "In progress",  "Naz, Marika",               "SES Directorate", "",                                "All of UnLtd"),
    ("Get legal advice",             "In progress",  "Beckie, Marika",             "SES Directorate", "",                                "All of UnLtd"),
    ("Events",                       "In progress",  "Keely, Asli, Rokaya, Fern",  "SES Directorate", "",                                "All of UnLtd"),
    ("Impact Measurement Toolkit",   "In progress",  "Tom, Alice",                 "SES Directorate", "",                                "All of UnLtd"),
    ("Your grant",                   "In progress",  "Fern",                       "SES Directorate", "",                                "All of UnLtd"),
    ("Media kit",                    "In progress",  "George, Fern",               "SES Directorate", "",                                "All of UnLtd"),
    ("Budget Guidelines",            "In progress",  "Fern",                       "SES Directorate", "",                                "All of UnLtd"),
    ("Safeguarding & child protection", "In progress", "Fern",                     "SES Directorate", "",                                "All of UnLtd"),
    ("SharePoint folder",            "In progress",  "Alice, Fern",                "SES Directorate", "",                                "All of UnLtd"),
    ("Onboarding form",              "In progress",  "Fern",                       "SES Directorate", "",                                "All of UnLtd"),
    ("End of award form",            "In progress",  "Alice, Fern",                "SES Directorate", "",                                "All of UnLtd"),
]

aw_table = doc.add_table(rows=1, cols=6)
aw_table.style = "Table Grid"
add_table_header_row(aw_table, aw_headers, col_widths_cm=[4.5, 2.5, 4.0, 3.5, 4.5, 2.5])
for i, row in enumerate(aw_rows):
    add_data_row(aw_table, row, alt=(i % 2 == 1))

spacer(doc, 10)

# ── About Awards space ────────────────────────────────────────────────────────
heading(doc, "About Awards Space", level=2, color=DARK_TEXT)
body(doc,
     "Public-facing. Subject to approval from Liam and Nicola in the context of the "
     "upcoming website rebrand. Structure and content may evolve.")

aa_headers = ["Page", "Status", "Consulted on first draft", "Editors", "Reviewer", "Commenters"]
aa_rows = [
    ("Welcome",                                    "Not yet migrated", "Fern",  "SES Directorate", "User Journey Support Lead (Fern)", "All of UnLtd"),
    ("Assistance for those living in Wales",       "Not yet migrated", "Fern",  "SES Directorate", "User Journey Support Lead (Fern)", "All of UnLtd"),
    ("Assistance with additional support needs",   "Not yet migrated", "",      "SES Directorate", "",                                "All of UnLtd"),
    ("Award Budget Guidelines",                    "Not yet migrated", "Fern",  "SES Directorate", "User Journey Support Lead (Fern)", "All of UnLtd"),
    ("Workshop Catalogue",                         "Not yet migrated", "",      "SES Directorate", "",                                "All of UnLtd"),
    ("How is my application assessed?",            "Not yet migrated", "",      "SES Directorate", "",                                "All of UnLtd"),
    ("Panel Meeting Guide",                        "Not yet migrated", "",      "SES Directorate", "",                                "All of UnLtd"),
    ("Panel Pages (dynamic)",                      "Not yet migrated", "",      "SES Directorate", "",                                "All of UnLtd"),
    ("Prepare your financial data",                "Not yet migrated", "",      "SES Directorate", "",                                "All of UnLtd"),
]

aa_table = doc.add_table(rows=1, cols=6)
aa_table.style = "Table Grid"
add_table_header_row(aa_table, aa_headers, col_widths_cm=[5.0, 3.0, 3.5, 3.5, 4.5, 2.5])
for i, row in enumerate(aa_rows):
    add_data_row(aa_table, row, alt=(i % 2 == 1))

spacer(doc, 10)

# ── Learn space ────────────────────────────────────────────────────────────────
heading(doc, "Learn Space", level=2, color=DARK_TEXT)
body(doc,
     "Public-facing. Will include Financial Fundamentals, publicly available UnLtd "
     "resources, and business-oriented resources from Tim. The plan is to have this space "
     "organised by July so SES colleagues can review, contribute to, and build out the "
     "Business Foundations section throughout August, ahead of the September public launch.")

body(doc,
     "[Tom / Joel to populate this table once the Learn space content list is confirmed]",
     )
p = doc.paragraphs[-1]
for run in p.runs:
    run.italic = True
    run.font.color.rgb = MID_GREY


# ─────────────────────────────────────────────────────────────────────────────
# 7. WHAT HAS MOVED TO GITBOOK / WHAT IS MOVING?
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "7.  What Has Already Moved to GitBook? What Is Moving?", level=1, color=UNLTD_RED)

body(doc,
     "[To be completed — provide a brief summary of resources already migrated and "
     "those still to come. This will help colleagues understand the current state ahead "
     "of the internal launch session.]")
p = doc.paragraphs[-1]
for run in p.runs:
    run.italic = True
    run.font.color.rgb = MID_GREY


# ─────────────────────────────────────────────────────────────────────────────
# 8. ROLES AND RESPONSIBILITIES
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "8.  Roles and Responsibilities", level=1, color=UNLTD_RED)

body(doc,
     "GitBook uses role-based permissions. Each role has a clearly defined scope of "
     "responsibility within the Hub:")

roles_headers = ["Role", "Who", "What they can do"]
roles_rows = [
    ("Hub Owner / Admin",  "Tom Sheppard",                             "Full access. Configures the Hub, manages permissions, publishes content, and oversees the overall quality and direction of the platform."),
    ("Editor",             "SES Directorate (role-based)",             "Can create, edit, and update content pages. Responsible for keeping content accurate and current within their area."),
    ("Reviewer",           "User Journey Support Lead (Fern) + others", "Reviews and signs off content before it is published. Ensures accuracy, tone, and alignment with UnLtd standards."),
    ("Commenter",          "All of UnLtd",                             "Can add comments and suggestions to content but cannot edit or publish. Provides feedback during review cycles."),
    ("Reader (Award Winner)", "Current award winners",                 "Can view gated Award Winner content. Invited by Fern as part of the welcome communication."),
]

roles_table = doc.add_table(rows=1, cols=3)
roles_table.style = "Table Grid"
add_table_header_row(roles_table, roles_headers, col_widths_cm=[4.0, 5.5, 12.5])
for i, row in enumerate(roles_rows):
    add_data_row(roles_table, row, alt=(i % 2 == 1))

spacer(doc, 8)
callout(doc,
        "Joel to confirm final Award Winner roles and permissions before the internal launch "
        "session. Tom will implement once agreed.",
        bg_hex="FFF8E1", prefix="⚠️  Action required: ")


# ─────────────────────────────────────────────────────────────────────────────
# 9. ROLLOUT TIMELINE
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "9.  Rollout Timeline", level=1, color=UNLTD_RED)

timeline_headers = ["Phase", "Target Date", "Audience", "Key Activities", "Actions Required"]
timeline_rows = [
    (
        "Phase 1: Internal Launch",
        "Early July 2026",
        "All UnLtd staff",
        (
            "• Showcase the Hub to the organisation\n"
            "• Onboard all staff onto GitBook\n"
            "• How-to session: making change requests and suggesting improvements\n"
            "• SESMs and UJSMs review content and provide structured feedback\n"
            "• Fern sends Hub link to new award winner cohort (replacing welcome pack)"
        ),
        (
            "• All staff: create your GitBook account and explore the Hub\n"
            "• SESMs and UJSMs: read content carefully and note gaps, errors, or suggestions\n"
            "• All: attend the onboarding session\n"
            "• All: submit feedback by [date TBC]\n"
            "• Fern: send Hub link in welcome communication to new cohort\n"
            "• SESMs and UJSMs: direct award winners to Hub resources during onboarding"
        ),
    ),
    (
        "Phase 2: Build-out",
        "August 2026",
        "SES Directorate, Delivery & Investment",
        (
            "• Review and build out the Learn — Business Foundations section\n"
            "• SESMs and UJSMs contribute resources and knowledge to the Learn space\n"
            "• All content changes completed and signed off ahead of September launch\n"
            "• Monthly practice sessions begin (see Ongoing Engagement)"
        ),
        (
            "• SESMs and UJSMs: continue submitting change requests as you use the Hub\n"
            "• Hub owner to share analytics and feedback summaries — please review and respond\n"
            "• All content to be signed off before public launch"
        ),
    ),
    (
        "Phase 3: Public Launch",
        "September 2026\n(aligned to new cohort)",
        "General public + award winners",
        (
            "• Hub made publicly accessible\n"
            "• Award Winner section remains gated\n"
            "• Full social entrepreneur user journey live\n"
            "• Fern sends Hub link in place of welcome pack\n"
            "• Award winners onboarded onto GitBook\n"
            "• Analytics reviewed; feedback gathered from first cohort"
        ),
        (
            "• Hub owner: publish public-facing content\n"
            "• Hub owner: decommission the beta hub on unltd.org.uk\n"
            "• Fern: invite new award winner cohort via GitBook\n"
            "• All: monitor and flag any issues reported by award winners"
        ),
    ),
]

timeline_table = doc.add_table(rows=1, cols=5)
timeline_table.style = "Table Grid"
add_table_header_row(timeline_table, timeline_headers, col_widths_cm=[4.0, 3.0, 3.5, 7.0, 7.0])
for i, row in enumerate(timeline_rows):
    add_data_row(timeline_table, row, alt=(i % 2 == 1))

spacer(doc, 10)


# ─────────────────────────────────────────────────────────────────────────────
# 10. ONGOING ENGAGEMENT
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "10.  Ongoing Engagement", level=1, color=UNLTD_RED)

body(doc,
     "To support adoption and hold SES and Delivery & Investment accountable to changing "
     "their processes, a regular practice session will be introduced alongside the Hub launch. "
     "These monthly sessions will bring the team together to:")

bullet(doc, "Ground GitBook use in day-to-day practice.")
bullet(doc, "Share knowledge across the team on topics such as finance, branding, and "
            "entrepreneurship.")
bullet(doc, "Discuss new resources or tools that could benefit social entrepreneurs.")
bullet(doc, "Build a shared culture of learning and continuous improvement.")

body(doc,
     "This model is informed by the Spark & Mettle Power Hour approach and already has "
     "buy-in from SESMs. Further details on format and cadence to be agreed.")


# ─────────────────────────────────────────────────────────────────────────────
# 11. COMMUNICATIONS AND TRAINING PLAN
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "11.  Communications and Training Plan", level=1, color=UNLTD_RED)

heading(doc, "Internal Launch Communication", level=2, color=DARK_TEXT)
body(doc,
     "Tom and Joel will draft a presentation and accompanying email to the whole "
     "organisation for the internal launch. This should:")

bullet(doc, "Demonstrate the Hub from the perspective of a social entrepreneur across "
            "their full lifecycle — discovery, application, and award winner.")
bullet(doc, "Show how a Support Manager and User Journey Support Manager would use GitBook "
            "to support an entrepreneur through their journey.")
bullet(doc, "Explain the change in process: GitBook links, not files.")

heading(doc, "Training", level=2, color=DARK_TEXT)
body(doc,
     "Tom will create training videos for each GitBook user role (reviewer, editor, "
     "commenter) before the internal launch. All SESMs and UJSMs will be trained on "
     "GitBook ahead of the public launch in September.")

heading(doc, "Award Winner Launch Communication", level=2, color=DARK_TEXT)
body(doc,
     "An email announcement will be sent to all award winners at the point of public "
     "launch. This will include a compelling value proposition for the Hub and either "
     "a live introductory workshop or a recorded video walkthrough.")


# ─────────────────────────────────────────────────────────────────────────────
# 12. OPEN QUESTIONS
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "12.  Open Questions", level=1, color=UNLTD_RED)

body(doc, "The following questions remain open and need to be resolved before or at the "
          "internal launch session:")

open_qs = [
    ("Award Winner roles and permissions", "Joel to review and confirm. Tom to implement once agreed."),
    ("Change request process", "What is the agreed process for SESMs and UJSMs to submit change requests to the Hub?"),
    ("Award winner onboarding flow", "How will award winners be invited onto GitBook? Confirmed: they will be invited after the launch meeting."),
    ("Analytics sharing", "How will analytics data be shared with the wider team? Built-in GitBook analytics will provide the baseline."),
    ("About Awards section", "Pending approval from Liam and Nicola in the context of the website rebrand. What is the timeline for that decision?"),
    ("Spark & Mettle", "What is the current status of the Spark & Mettle partnership, and are there implications for the Hub content or audience?"),
    ("Training format", "Will Award Winner training take the form of a live workshop, recorded video, or written guide?"),
    ("Hub branding", "What will the Award Winner Hub be called? How should it be branded?"),
    ("Enforcement of new process", "How will we handle SESMs who continue to share files from their own devices after launch?"),
    ("Tom's attendance at team meetings", "Tom to attend a Friday end-of-week session (per Tim) to introduce the project — format and timing to be confirmed."),
]

oq_table = doc.add_table(rows=1, cols=3)
oq_table.style = "Table Grid"
add_table_header_row(oq_table, ["Question", "Owner", "Notes / Status"], col_widths_cm=[6.5, 3.0, 12.5])
for i, (question, note) in enumerate(open_qs):
    owner = "Joel" if "Joel" in note else ("Tom" if "Tom" in note else "TBC")
    add_data_row(oq_table, [question, owner, note], alt=(i % 2 == 1))

spacer(doc, 10)


# ─────────────────────────────────────────────────────────────────────────────
# 13. NEXT STEPS AND DECISIONS
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "13.  Next Steps and Decisions", level=1, color=UNLTD_RED)

body(doc, "The following actions and decisions are required to keep the project on track:")

actions_headers = ["Action", "Owner", "Due", "Status"]
actions_rows = [
    ("Confirm Award Winner roles and permissions in GitBook",                  "Joel",        "Before internal launch", "Open"),
    ("Implement roles and permissions in GitBook",                             "Tom",         "Once Joel confirms",     "Open"),
    ("Create GitBook training videos (reviewer, editor, commenter)",          "Tom",         "Before internal launch", "Open"),
    ("Draft internal launch presentation and org-wide email",                 "Tom + Joel",  "Before internal launch", "Open"),
    ("Agree change request process for SESMs and UJSMs",                      "Tom + Joel",  "Before internal launch", "Open"),
    ("Confirm About Awards section with Liam and Nicola",                     "Tom / Fern",  "Before July",            "Open"),
    ("Plan and schedule monthly practice sessions",                           "Tom + Joel",  "Before internal launch", "Open"),
    ("Confirm Tom attendance at Friday team session to introduce the project","Tom / Tim",   "TBC",                    "Open"),
    ("Decommission beta hub on unltd.org.uk",                                 "Tom",         "September launch",       "Open"),
    ("Send launch email to all award winners",                                "Fern",        "September launch",       "Open"),
]

actions_table = doc.add_table(rows=1, cols=4)
actions_table.style = "Table Grid"
add_table_header_row(actions_table, actions_headers, col_widths_cm=[9.0, 3.0, 4.5, 2.5])
for i, row in enumerate(actions_rows):
    add_data_row(actions_table, row, alt=(i % 2 == 1))

spacer(doc, 10)


# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
output_path = "/Users/tom/Documents/GitHub/unltd/UnLtd Hub — Internal Engagement Plan.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
